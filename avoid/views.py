from django.shortcuts import render
import docx2txt

# Create your views here.
from django.shortcuts import render, redirect
from django.http import HttpResponse
import language_tool_python
import json
import requests
import os,re
import docx
import io
from docxlatex import Document
import pickle as pkl
import subprocess
import spacy
import time
from urllib.parse import urlparse
from docx2python import docx2python
import hashlib
import asyncio
from reports.models import Report
from django.core.files import File
from django.core.files.storage import FileSystemStorage
from datetime import datetime
from stats.models import Stats
import shutil
from urllib.parse import quote
from django.utils.timezone import localtime, now
from django.contrib.auth.decorators import login_required
from django.core.files.storage import FileSystemStorage
import gzip, io

USER_FOLDER = ''
ROOT_FOR_USER_FOLDERS = 'files/user_folders/'


def generate_graph(message_count:dict):
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('agg')

    # Define a color palette with shades of blue
    colors = plt.cm.Blues_r(range(len(message_count)))

    y_labels = [msg for msg in message_count.keys()]
    values = [val for val in message_count.values()]

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.barh(y_labels, values, height=0.8, color=colors)

    # Add grid lines to the chart
    ax.grid(axis='x', linestyle='--', alpha=0.7)

    # Remove axes and horizontal tick marks
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.tick_params(axis='both', which='both', length=0, labelsize=12)
    ax.set_xticks([])

    # Rotate the y-axis labels to vertical orientation
    ax.set_yticklabels(y_labels, rotation=0, fontsize=20)

    # Add values to bars
    for i, v in enumerate(values):
        ax.text(v + 0.5, i, str(v), color='black', fontweight='bold', fontsize=18)

    # Save the chart as a high-resolution image
    plt.savefig('files/graph.png', bbox_inches='tight', pad_inches=0)

    # Close the figure to free up memory
    plt.close()

def name_the_file(username, type):
    date = datetime.now().strftime("%d.%m.%y")
    # client_time = localtime(now()).strftime("%d.%m.%y")
    with open(USER_FOLDER + "file_name.pkl", "rb") as f:
        doc_name = pkl.load(f)
    file_name = ''.join(doc_name.split('.')[:-1]) + ' — '
    if type == 'spravka':
        file_name += 'Справка'
    elif type == 'avoid_report':
        file_name += 'Проверка на попытки обхода'
    elif type == 'plagiarism_report':
        file_name += 'Проверка на плагиат'
    file_name += ' — ' + date
    return file_name

def generate_report_word(report_data, white_symbols: list,  symbols_count, author, text_formulas_russian, user, topic, file_name):
    from docxtpl import DocxTemplate

    doc = DocxTemplate('files/report.docx')
    # !! context should be derived from single source
    context = {'author_name': ','.join(author), 'topic': topic, 'symbols_count':symbols_count, 'work_checker': f'{user.first_name} {user.last_name}'}
    doc.render(context)

    # generate report table with messages
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Тип ошибки'
    hdr_cells[1].text = 'Контекст ошибки'

    for i, row in enumerate(report_data):
        cells = table.add_row().cells
        cells[0].text = row['message']
        context_cell = cells[1].paragraphs[0]
        context_text = row['context']
        context_cell.add_run( context_text[:row['offset']] )
        context_cell.add_run( context_text[row['offset']:row['offset']+row['errorLength']]).font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
        context_cell.add_run( context_text[row['offset'] + row['errorLength']:] )

    for row in white_symbols:
        cells = table.add_row().cells
        cells[0].text = 'Попытка скрыть символы'
        context_cell = cells[1].paragraphs[0]
        context_text = row['context']
        context_cell.add_run( context_text[:row['offset']] )
        context_cell.add_run( context_text[row['offset']:row['offset']+row['errorLength']]).font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
        context_cell.add_run( context_text[row['offset'] + row['errorLength']:] )
    
    for formula in text_formulas_russian:
        cells = table.add_row().cells
        cells[0].text = 'Попытка заменить текст формулой'
        context_cell = cells[1].paragraphs[0]
        context_text = formula['context']
        context_cell.add_run( context_text[:formula['offset']] )
        context_cell.add_run( context_text[formula['offset']:formula['offset']+formula['errorLength']]).font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
        context_cell.add_run( context_text[formula['offset'] + formula['errorLength']:] )

    #output = io.BytesIO()
    input_docx = USER_FOLDER + 'report.docx'
    doc.save(input_docx)
    
    avoid_report_name = name_the_file(user.username, 'avoid_report')
    date_time = datetime.now().strftime("%d.%m.%y %H:%M")
    # date_time = localtime(now()).strftime("%d.%m.%y %H:%M")
    report = Report(name=avoid_report_name, user=user, date_str=date_time)
    # with open(f"files/{user.username}/file_name.pkl", "rb") as f:
    #     file_name = pkl.load(f)
    with open(USER_FOLDER + 'report.docx', 'rb') as f:
        report_file = File(f)
        report.file.save(avoid_report_name + '.docx', report_file)
        report_file.close()

def generate_file_word(message_count, authors, user, topic):
    from docxtpl import DocxTemplate
    doc = DocxTemplate('files/spravka.docx')
    # context = { dynamic data derived from document parsing }
    context = {'author_name': ','.join(authors), 'topic': topic,'work_checker': f'{user.first_name} {user.last_name}'}
    doc.render(context)
    generate_graph(message_count)
    graph_size = 8 if len(message_count)<2 else 15
    doc.add_picture('files/graph.png', docx.shared.Cm(graph_size))

    doc.add_paragraph('Работу проверил: ' + context['work_checker'])
    date_paragraph = doc.add_paragraph('Дата подписи: _________   ')
    date_paragraph.add_run("\t"*3)
    run = date_paragraph.add_run("_" * 45)
    run.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
    signature_paragraph = doc.add_paragraph()
    run = signature_paragraph.add_run("Подпись Проверяющего")
    run.font.size = docx.shared.Pt(10)
    signature_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

    # output = io.BytesIO()
    input_docx = USER_FOLDER + 'spravka.docx'
    doc.save(input_docx)
    # Rewind the file pointer to the beginning of the file
    #output.seek(0)
    subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", input_docx, "--outdir", USER_FOLDER])
    
def get_text_from_first_page(filepath):
    doc = docx.Document(filepath)
    text = ""
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading') or para.style.name == 'Title':
            # Stop when a heading or title style is encountered
            break
        text += para.text + " "
    return text.strip()

def find_topic(text):
    
    text = text.lower()


    match = re.search(r'(тема:|по теме:).*?([^\w\s]).*?', text, re.DOTALL)

    if match:
        topic_text = match.group(0)
        # Remove the keyword and punctuation from the topic text
        topic_text = topic_text.split(':', 1)[-1].strip()
        topic_text = re.sub(r'[^\w\s]$','',topic_text)
        print(topic_text)
        return topic_text


    return "не найден"

def find_capitalized_authors(text):
    nlp = spacy.load('ru_core_news_sm')
    doc = nlp(text)

    authors = []
    for ent in doc.ents:
        if ent.label_ == 'PER' and ent.text.istitle():
            authors.append(ent.text)


    return authors

def clean_old_user_folders():
    current_time = time.time()
    threshold_time = 86400     # 86400s = 24h
    for folder in os.listdir(ROOT_FOR_USER_FOLDERS):
        folder_path = os.path.join(ROOT_FOR_USER_FOLDERS, folder)
        if os.path.isdir(folder_path):
            folder_mtime = os.path.getmtime(folder_path)
            if current_time - folder_mtime > threshold_time:
                shutil.rmtree(folder_path)

def create_stats_for_old_users():
    from django.contrib.auth.models import User
    from stats.models import Stats
    for user in User.objects.all():
        stats = Stats.objects.filter(user=user).first()
        if not stats:
            stats = Stats(user=user)
            stats.reports_count = 0
            stats.symbols_used = 0
            stats.symbols_left = 1000000
            stats.save()

# def unpack_gz_file(request):
#     USER_FOLDER = f'files/user_folders/{request.user.username}/'
#     report = Report.objects.get(id=471) # 471 – строка для теста
#     gz_file = report.pickles
#     with gzip.open(gz_file, "rb") as f:
#         pickles = pkl.load(f)
#         for filename, data in pickles.items():
#             output_filename = USER_FOLDER + 'temp/' + filename
#             with open(output_filename, "wb") as f:
#                 pkl.dump(data, f)

def merge_and_compress_pkl_files(request):
    pickles = {}
    USER_FOLDER = f'files/user_folders/{request.user.username}/'
    for filename in os.listdir(USER_FOLDER):
        if filename.endswith(".pkl"):
            with open(os.path.join(USER_FOLDER, filename), "rb") as f:
                data = pkl.load(f)
                pickles[filename] = data
    # output_filename = os.path.join(USER_FOLDER, "pickles.gz")
    # with gzip.open(output_filename, "wb") as f:
    #     pkl.dump(pickles, f)
    # return output_filename
    output_file = io.BytesIO()
    with gzip.open(output_file, "wb") as f:
        pkl.dump(pickles, f)
    return output_file

def test_func(request):
    megapkl = merge_and_compress_pkl_files(request)
    # filename = name_the_file(username, 'plagiarism_report')
    USER_FOLDER = f'files/user_folders/{request.user.username}/'
    date_time = datetime.now().strftime("%d.%m.%y %H:%M")
    with open(USER_FOLDER + 'file_name.pkl', 'rb') as f:
        file_name = pkl.load(f)
    file_name = ''.join(file_name.split('.')[:-1])
    report = Report(name=file_name, user=request.user, date_str=date_time)
    with open(USER_FOLDER + file_name + '.docx', 'rb') as f:
        original_doc = File(f)
        report.original_doc.save(file_name + '.docx', original_doc)
    # with open(USER_FOLDER + 'plagiarism_report.docx', 'rb') as f:
    #     report_file = File(f)
    #     report.report.save(filename + '.docx', report_file)
    report.pickles.save('pickles.gz', megapkl)
    report.save()
    # unpack_gz_file(request)

@login_required
def file_upload_avoid(request):
    # test_func(request)
    # date = request.POST.get('date')
    # print(date)
    # clean_old_user_folders()            # delete user folders for user temp files, if folders older than an hour
    create_stats_for_old_users()
    if request.method == 'POST':
        avoid_file = request.FILES['avoid_file']
        avoid_file_name = avoid_file.name
        suff = ('.doc', '.docx')
        report = []
        white_symbols = []
        symbols = []
        response_context = {}
        username = request.user.username
        global USER_FOLDER
        USER_FOLDER = ROOT_FOR_USER_FOLDERS + username + '/'
                # save the avoid file to the server
        
        if avoid_file_name.lower().endswith(suff):
            #
            with open(USER_FOLDER + avoid_file_name, 'wb') as f:
                f.write(avoid_file.read())
            txt_avoid_filename = ''.join(avoid_file_name.split('.')[:-1])+'.txt'
            file_path = USER_FOLDER + avoid_file_name
            if not os.path.exists(USER_FOLDER):
                os.mkdir(USER_FOLDER)
            if avoid_file_name.lower().endswith('.doc'):
                with open(file_path, 'ab') as docf:
                    for ch in avoid_file.chunks():
                        docf.write(ch)
                avoid_file_name_docx = avoid_file_name + 'x'
                subprocess.call(
                    ["soffice", "--headless", "--convert-to", "docx", file_path, "--outdir", USER_FOLDER])
                file_path = USER_FOLDER + avoid_file_name_docx
        elif avoid_file_name.lower().endswith('.pdf'):
            if not os.path.exists(USER_FOLDER):
                os.mkdir(USER_FOLDER)
            FileSystemStorage(location=USER_FOLDER).save(avoid_file_name.lower(), avoid_file)
            subprocess.call(["soffice","--infilter=writer_pdf_import","--convert-to","docx","--outdir",USER_FOLDER,USER_FOLDER+avoid_file_name.lower()])
            avoid_file_name = ''.join(avoid_file_name.lower().split('.')[:-1])+'.docx'
            avoid_file = open(USER_FOLDER+avoid_file_name, 'rb')
        else:
            return "error"
        first_page_text = get_text_from_first_page(avoid_file)
        topic = find_topic(first_page_text)
        with open(USER_FOLDER + 'topic.pkl', 'wb') as f:
            pkl.dump(topic, f)
        authors = find_capitalized_authors(first_page_text)
        with open(USER_FOLDER + 'authors.pkl','wb') as f:

            pkl.dump(authors, f)
        # save the avoid file to the server
        if avoid_file_name.lower().endswith(suff):

            #
            txt_avoid_filename = ''.join(avoid_file_name.split('.')[:-1])+'.txt'
            if avoid_file_name.lower().endswith('.doc'):
                with open(USER_FOLDER + avoid_file_name,'ab') as docf:
                    for ch in avoid_file.chunks():
                        docf.write(ch)
                avoid_file_name_docx = avoid_file_name + 'x'
                subprocess.call(["soffice", "--headless", "--convert-to", "docx", USER_FOLDER + avoid_file_name, "--outdir", USER_FOLDER])
                doc = docx.Document(USER_FOLDER + avoid_file_name_docx)
            else:
                doc = docx.Document(avoid_file)

            symbol_count = 0
            for para in doc.paragraphs:
                symbol_count += len(para.text)

            with open(USER_FOLDER + 'original_symbol_count.pkl', 'wb') as f:
                pkl.dump(symbol_count, f)

            print(symbol_count)
            # fullText = []
            # start processing white symbols
            for i, para in enumerate(doc.paragraphs):
                for j, run in enumerate(para.runs):
                    if run.font.color.rgb == docx.shared.RGBColor(255, 255, 255):
                        # locate and save white symbol for further processing
                        white_symbol_report = {}
                        symbols.append(run.text)
                        key = f'Белый символ {i}.{j}'  # Use paragraph and run index in key
                        # Find the previous and next words
                        #f'Параграф {i}, Символ {run.text}'
                        # words = para.text.split()
                        prev_word = ''.join(list(map(lambda run: run.text, para.runs[j-5:j])))
                        next_word = ''.join(list(map(lambda run: run.text, para.runs[j+1:j+5]))) if j < len(para.runs)-1 else ''
                        white_symbol_report['offset'] = len(prev_word)
                        symbol_text = prev_word + run.text + next_word
                        white_symbol_report['context'] = symbol_text
                        white_symbol_report['errorLength'] = len(run.text)
                        white_symbols.append(white_symbol_report)
                #fullText.append(para.text)
            #fullText = '\n'.join(fullText)
            pattern = re.compile("[а-яА-ЯёЁ]+")
            doc_latex = Document(avoid_file, inline_delimiter='$ineq$') # ineq = inline equation
            text_with_formulas = doc_latex.get_text()
            # count number of formulas
            # formulas_count = text_with_formulas.count('$ineq$')//2
            # text_with_formulas = text_with_formulas.replace('$ineq$', '')
            formulas_count = 0
            text_formulas_with_russian = []
            for formula in re.finditer(r'\$ineq\$.+?\$ineq\$', text_with_formulas):
                match = re.search(pattern, text_with_formulas[formula.start() : formula.end()])
                if match:
                    formulas_count += 1
                    formula_report = {}
                    formula_text = match.group()
                    context_start = max(formula.start()-20, 0)
                    context_end = min( formula.end() + 20, len(text_with_formulas) )
                    # context = '... $ineq$ русские буквы $ineq$ ...'
                    context = text_with_formulas[context_start:formula.start()] + formula_text + text_with_formulas[formula.end():context_end]
                    formula_report['context'] = context
                    formula_report['errorLength'] = match.end() - match.start()
                    formula_report['offset'] = formula.start() - context_start
                    text_formulas_with_russian.append(formula_report)
                #print("Russian formula", formulas_count, ":", formula)
            #with open('files/' + )
            with open(USER_FOLDER + 'file_name.pkl','wb') as f:
                pkl.dump(txt_avoid_filename,f)
            
            text_with_formulas = re.sub(r'\$ineq\$ (.+?) \$ineq\$',r'\1',text_with_formulas)
            symbol_count = len(text_with_formulas)
            with open(USER_FOLDER + txt_avoid_filename, 'w+', encoding='UTF-8') as destination:
                destination.write(text_with_formulas)
            

                
        elif avoid_file_name.lower().endswith('.txt'):
            with open(USER_FOLDER + avoid_file_name, 'wb+') as destination:
                for chunk in avoid_file.chunks():
                    destination.write(chunk)
        else:
            return HttpResponse('Неверный формат файла')
                
        avoid_file_name = ''.join(avoid_file_name.split('.')[:-1])
        
        API_ENDPOINT = 'http://localhost:30000/avd/{}'.format(avoid_file_name)

        response = requests.get(API_ENDPOINT)

        if response.status_code == 200:
            # return redirect('avoid:file_processed_avoid', response=response.text)
            errors_list = ['WHITESPACE_RULE', 'Latin_single_letter', 'Latin_letters']
            data = json.loads(response.text)
            percentage = dict()
            messages_count = {}
            # data = [{'ruleId':'WHITESPACE_RULE' | 'MORFOLOGIK_RULE_RU_RU' | 'Latin_single_letter',
            #  'message':str, 'context':str(original text), ... }]
            # messages_count = map<'message_from_match', int>, e.g. {'some_message': int}
            kazcharacter_set = ['ә', 'ң', 'ө', 'һ', 'і', 'ғ', 'қ', 'ұ', 'ү']
            for match in data:
                if match['ruleId'] in errors_list:
                    if match['ruleId'] == 'Latin_single_letter' or match['ruleId'] == 'Latin_letters':
                        if any([kazchar in match['context'] for kazchar in kazcharacter_set]):
                            continue
                        match['message'] = 'Возможно нужна буква из кириллицы вместо аналогичной по начертанию латинской или наоборот'

                    report_cell = {
                                'message': match['message'],
                                'ruleId': match['ruleId'],
                                'context': match['context'],
                                'errorLength' : match['errorLength'],
                                'offset': match['offsetInContext'] if 'offsetInContext' in match else None
                            }
                    report.append(report_cell)

                    if match['message'] not in messages_count:
                        messages_count[match['message']] = 1
                    else:
                        messages_count[match['message']] += 1
            if formulas_count:
                messages_count['Потенциальная попытка заменить текст формулой'] = formulas_count
            if white_symbols:
                messages_count['Попытка скрыть символы'] = len(white_symbols)

            percentage = json.dumps(messages_count)
            response_context['percentage'] = percentage
            response_context['messages_count'] = messages_count
            response_context['len_text'] = symbol_count
            # save.pkl = Array<report_cell>
            with open(USER_FOLDER + 'save.pkl', 'wb') as f:
                pkl.dump(report, f)
            # message_count.pkl = map<'message_from_match', int>, messages from report + additional messages(uniqueness, white_symbols and etc.)
            with open(USER_FOLDER + 'message_count.pkl', 'wb') as f:
                pkl.dump(messages_count, f)
            with open(USER_FOLDER + 'white_symbols_report.pkl','wb') as f:
                pkl.dump(white_symbols,f)
            with open(USER_FOLDER + 'text_formulas_russian.pkl','wb') as f:
                pkl.dump(text_formulas_with_russian,f)
            with open(USER_FOLDER + 'symbols_count.pkl', 'wb') as f:
                pkl.dump(symbol_count, f)
            user = request.user
            generate_report_word(report, white_symbols, symbol_count,authors, text_formulas_with_russian, user, topic, avoid_file_name)
            if (messages_count.get('Попытка скрыть символы', 0) >= 1 or
        messages_count.get('Возможно нужна буква из кириллицы вместо аналогичной по начертанию латинской или наоборот', 0) >= 3 or
        messages_count.get('Количество формул',0) > 5 or
        messages_count.get('Попытка скрыть символы',0) >= 1 or
        messages_count.get('Повтор пробела', 0) > 15):
                response_context['first_result'] = 'К сожалению, Вы не прошли первый этап проверки! Вы можете увидеть подробности, скачав отчёт о попытках обхода.'
            else :
                response_context['first_result'] = 'Первый этап пройден успешно!'
            return render(request, 'avoid/file_processed_avoid.html', response_context)
        else:
            # the request was not successful, so handle the error here
            return HttpResponse('Извиняюсь, у вас ошибка, да....')
    return render(request, 'avoid/file_upload_avoid.html')

def avoid_func(av_text):
    my_tool = language_tool_python.LanguageTool('ru-RU')  
    
    text = av_text + ".txt"
    with open(f'files/{text}',encoding='UTF-8') as f:
        final_text = f.read()

    my_matches = my_tool.check(final_text)
    return my_matches

def plagiarism_uniqueness(request):
    username = request.user.username
    with open(USER_FOLDER + 'file_name.pkl','rb') as f:
        file_name = pkl.load(f)
    with open(USER_FOLDER + file_name, 'r', encoding = 'UTF-8') as f:
        final_text = f.read()
    json_final = plagiarism_check(final_text)
    if json_final.get('error'):
        return HttpResponse( json.dumps( {'error':json_final.get('error')} ) )
    result_json = json.loads(json_final['result_json'])
    with open(USER_FOLDER + 'json_final.pkl','wb') as f:
        pkl.dump(result_json,f)
    return HttpResponse(json.dumps({'uniqueness':json_final.get('unique'),'urls':result_json['urls']}))

def generate_plag_report(user):
    from docxtpl import DocxTemplate

    username = user.username
    doc = DocxTemplate('files/plagiarism_report.docx')
    # !! context should be derived from single source
    with open(USER_FOLDER + 'authors.pkl','rb') as f:
        authors = pkl.load(f)
    with open(USER_FOLDER + 'json_final.pkl','rb') as f:
        json_final = pkl.load(f)
    with open(USER_FOLDER + 'topic.pkl', 'rb') as f:
        topic = pkl.load(f)
    context = {
        'author_name': ','.join(authors),
        'topic': topic,
        'text_uniqueness': json_final['text_unique'],
        'work_checker': f'{user.first_name} {user.last_name}',
    }
    doc.render(context)

    # generate report table with messages
    table = doc.add_table(rows=1, cols=3)
    row_counter = 1
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '%'
    hdr_cells[1].text = 'Текст'
    hdr_cells[2].text = 'Ресурс заимствования'
    if type(json_final['result_json']) == str:
        result_json = json.loads(json_final['result_json'])
    else:
        result_json = json_final['result_json']
    text_words = result_json['clear_text'].split()
    
    for row in result_json['urls']:
        word_indices = list(map(int, row['words'].split()))
        cells = table.add_row().cells
        cells[0].text = str(row['plagiat'])
        borrowed_text = ''
        for idx in word_indices:
            borrowed_text = borrowed_text +' '+ text_words[idx]
        cells[1].text = borrowed_text
        cells[2].text = row['url']
        row_counter = row_counter + 1
    
    input_docx = USER_FOLDER + 'plagiarism_report.docx'
    doc.save(input_docx)

    plag_report_name = name_the_file(username, 'plagiarism_report')
    date_time = datetime.now().strftime("%d.%m.%y %H:%M")
    # date_time = localtime(now()).strftime("%d.%m.%y %H:%M")
    report = Report(name=plag_report_name, user=user, date_str=date_time)
    with open(USER_FOLDER + 'plagiarism_report.docx', 'rb') as f:
        report_file = File(f)
        report.file.save(plag_report_name + '.docx', report_file)
        report_file.close()

def get_plagiarism_report(request):
    username = request.user.username
    generate_full_text_report(username)
    
    # megapkl = create_megapkl()
    # # filename = name_the_file(username, 'plagiarism_report')
    # date_time = datetime.now().strftime("%d.%m.%y %H:%M")
    # with open('files/file_name.pkl', 'rb') as f:
    #     file_name = pkl.load(f)
    # report = Report.create(name=file_name, user=request.user, date_str=date_time)
    # with open(USER_FOLDER + file_name + '.txt', 'rb') as f:
    #     original_doc = File(f)
    #     report.original_doc.save(file_name + '.docx', original_doc)
    # # with open(USER_FOLDER + 'plagiarism_report.docx', 'rb') as f:
    # #     report_file = File(f)
    # #     report.report.save(filename + '.docx', report_file)
    # # report.pickles.save('megapkl.pkl.gz', megapkl)
    # # report.save()

    subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", USER_FOLDER + 'full_text_report_result.docx', "--outdir", USER_FOLDER])

    with open(USER_FOLDER + 'full_text_report_result.pdf', 'rb') as f:
        file_data = f.read()
    # Set the appropriate response headers to indicate a file download
    filename = name_the_file(username, 'plagiarism_report')
    quoted_filename = quote(filename.encode('utf-8'))
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{quoted_filename}.pdf'
    
    # Copy the contents of the file object to the response object
    response.write(file_data)
    
    return response

def check_available_symbols(user):
    username = user.username
    with open(USER_FOLDER + 'original_symbol_count.pkl','rb') as f:
        symbols_count = pkl.load(f)
    last_stats = Stats.objects.filter(user=user).last()
    symbols_left = last_stats.symbols_left
    # if symbols_left < 1000000:
    if symbols_left < symbols_count:
        return False
    reports_count = last_stats.reports_count
    symbols_used = last_stats.symbols_used
    symbols_left -= symbols_count
    symbols_used += symbols_count
    Stats.objects.create(user=user, reports_count=reports_count, symbols_used=symbols_used, symbols_left=symbols_left)
    return True

def plagiarism_check(final_text, user, excepturl=[]):
    username = user.username
    text_hash_filename= hashlib.sha256(final_text.encode()).hexdigest()
    if excepturl:   #for the secondary check ("Exclude" button \ кнопка "Исключить")
        with open(USER_FOLDER + text_hash_filename + '.json', 'r') as file:
            json_response_final = json.load(file)

        data = json_response_final['result_json']
        if type(data) == str:
            data=json.loads(data)
        clear_text = data['clear_text']
        all_word_count = len(clear_text.split())
        urls = data['urls']
        domains = []
        words_plag=[False]*all_word_count
        plag_dict={}
        filtered_urls = []
        for url in urls:
            domain = url['url'].split('/')[2]
            domains.append(domain)
            if domain in plag_dict:
                plag_dict[domain]=plag_dict[domain]+url['words'].split(' ')
            else:
                plag_dict[domain]=url['words'].split(' ')
            if domain not in excepturl:
                filtered_urls.append(url)
        for domain in plag_dict:
            if domain in excepturl:
                continue
            else:
                for word_index_str in plag_dict[domain]:
                    word_index = int(word_index_str)-1
                    words_plag[word_index] = True
        count_false = words_plag.count(False)
        data['urls'] = filtered_urls
        uniqueness=round(((count_false / all_word_count)) * 100, 2)
        data['unique'] = uniqueness
        json_response_final['result_json'] = data
        json_response_final['unique'] = uniqueness
        json_response_final['text_unique'] = uniqueness
        #print(uniqueness,"%")
        return json_response_final
    else:   # for the initial check ("Check for uniqueness" button \ кнопка "Проверить на уникальность")
        #################################################################
        from django.contrib import messages
    
        result = check_available_symbols(user)
        if not result:
            pass
            # messages.error(request, 'У вас недостаточно символов')
            # json_response_final = {'error': json_response_final.get('error_desc', 'Возможно закончилась подписка')}
            # return json_response_final
            # return HttpResponse("У вас недостаточно токено")
        
        #################################################################
        if BigFile.is_bigfile(final_text):
            json_response_final = BigFile.plagiarism_check(final_text)
        else:
            url = 'https://api.text.ru/post'
            params = {
                'userkey': '31b045c683167166b5bae6c1014561cd',
                'text': final_text,
                'excepturl': ','.join(excepturl)
            }
            response = requests.post(url, data=params)
            json_response = response.json()
            if response.status_code == 200 and json_response.get('error_code') is None:
                time.sleep(60)
                url = 'https://api.text.ru/post'
                params = {
                    'userkey': '31b045c683167166b5bae6c1014561cd',
                    'uid': json_response['text_uid'],
                    'jsonvisible': 'detail'
                }
                response = requests.post(url, data=params)
                json_response_final = response.json()
                json_num = json_response_final.get('error_code')
                #print(json_response_final, json_num)
                if (response.status_code == 200 and json_response_final.get('error_code') is None) :
                    while json_num == 181:
                        time.sleep(10)
                        response = requests.post(url, data=params)
                        json_response_final = response.json()
                        json_num = json_response_final.get('error_code')
                        #print(json_response_final, json_num)
                else:
                    json_response_final = {'error': json_response_final.get('error_desc', 'Возможно закончилась подписка')}

            else:
                json_response_final = {'error': json_response.get('error_desc', 'Error occurred while processing the file')}

        with open(USER_FOLDER + text_hash_filename + '.json', 'w') as file:
            # serialize the dictionary to a JSON string and write it to the file
            json.dump(json_response_final, file)
        return json_response_final

def generate_full_text_report(username):
    def split_to_subtokens(str):
        word_token = ""
        subtokens = []
        for c in str:
            if c.isalpha():
                word_token += c
            elif c.isnumeric():
                word_token += c
            else:
                if c=="'" and word_token:
                    continue
                if c=='№':
                    c='No'
                    word_token += c
                    continue
                if word_token:
                    subtokens.append(word_token)
                    word_token=""
                subtokens.append(c)
        if word_token:
            subtokens.append(word_token)
        return subtokens

    with open('files/json_final.pkl','rb') as f:
        json_final = pkl.load(f)
    if type(json_final['result_json']) == str:
        result_json = json.loads(json_final['result_json'])
    else:
        result_json = json_final['result_json']
    with open('files/file_name.pkl', 'rb') as f:
        file_name = pkl.load(f)
    with open('files/' + file_name, 'r', encoding='UTF-8') as f:
        orig_text = f.read()
    orig_txt_tokenswithchars = orig_text.split()
    clear_txt_tokens = result_json['clear_text'].split()
    clear_text_ind = 0
    orig_text_ind = 0
    mapping = {}
    final_tokens=[]
    final_tokens_idx=0

    try:
        while clear_text_ind < len(clear_txt_tokens):
            if orig_txt_tokenswithchars[orig_text_ind] == clear_txt_tokens[clear_text_ind]:
                final_tokens.append(orig_txt_tokenswithchars[orig_text_ind])
                final_tokens.append(' ')
                mapping[clear_text_ind] = final_tokens_idx
                clear_text_ind +=1
                orig_text_ind +=1
                final_tokens_idx +=2
            else:
                token = orig_txt_tokenswithchars[orig_text_ind]
                subtokens = split_to_subtokens(token)
                for st in subtokens:
                    if clear_text_ind == len(clear_txt_tokens):
                        final_tokens.append(st)
                        final_tokens_idx += 1
                        continue
                    if st == clear_txt_tokens[clear_text_ind]:
                        mapping[clear_text_ind] = final_tokens_idx
                        clear_text_ind +=1
                        final_tokens_idx += 1
                        final_tokens.append(st)
                    else:
                        final_tokens.append(st)
                        final_tokens_idx += 1
                orig_text_ind +=1
                final_tokens.append(' ')
                final_tokens_idx += 1
    except:
        mapping = {i:i for i in range(len(clear_txt_tokens))}
        final_tokens=[c+' ' for c in clear_txt_tokens]

# 1 create common base for colored words e.g. {1:red, 2:blue, 3: green, etc.}
# 2 create mapping from colors to urls e.g. {'yellow':url1, 'red':url2, etc.}
    colors = [[255, 0, 0],       # Red
                [0, 255, 0],     # Green
                [0, 0, 255],     # Blue
                [255, 255, 0],   # Yellow
                [255, 0, 255],   # Magenta
                [0, 255, 255],   # Cyan
                [255, 128, 0],   # Orange
                [128, 0, 255],   # Purple
                [0, 255, 128],   # Lime
                [128, 255, 0],   # Chartreuse
                [255, 0, 128],   # Rose
                [0, 128, 255],   # Azure
                [255, 128, 128], # Coral
                [128, 255, 128], # Spring Green
                [128, 128, 255], # Violet
                [255, 255, 128], # Pastel Yellow
                [255, 128, 255], # Pastel Magenta
                [128, 255, 255], # Pastel Cyan
                [192, 192, 192], # Silver
                [128, 128, 128], # Gray
                [0, 0, 128],     # Navy
                [0, 128, 128],   # Teal
                [128, 0, 0],     # Maroon
                [128, 0, 128],   # Purple
                [0, 128, 0],     # Olive
                [128, 128, 0],   # Olive Green
                [139, 69, 19],   # Saddle Brown
                [139, 0, 139],   # Dark Magenta
                [0, 139, 139],   # Dark Cyan
                [139, 0, 0],     # Dark Red
                [0, 139, 0],     # Dark Green
                [0, 0, 139],     # Dark Blue
                [218, 165, 32],  # Goldenrod
                [139, 69, 19],   # Chocolate
                [205, 133, 63],  # Peru
                [0, 0, 0],       # Black
                [255, 255, 255], # White
                [128, 0, 0],     # Dark Maroon
                [0, 128, 0],     # Dark Olive Green
                [0, 0, 128]]
    colored_words = {}
    url2color={}
    next_color = 0
    for url in result_json['urls']:
        words = list(map(int,url['words'].split()))
        words_transformed = [mapping[word] for word in words]
        included = False
        urlcolor = docx.shared.RGBColor(*colors[next_color])
        for word in words_transformed:
            if word not in colored_words:
                colored_words[word] = urlcolor
                included = True
        if included:
            next_color = next_color+1 if next_color<(len(colors)-1) else 0
            url2color[url['url']] = urlcolor
    
    doc = docx.Document('files/full_text_report.docx')
    doc.add_paragraph("Процент уникальности: " + str(result_json['unique']))
    par = doc.add_paragraph()
    from collections import OrderedDict
    colored_words = OrderedDict(sorted(colored_words.items()))
    counter=0

    for word,color in colored_words.items():
        par.add_run(final_tokens[counter:word])
        par.add_run(final_tokens[word]).font.color.rgb = color
        counter = word + 1
    par.add_run(final_tokens[counter:])
    
    doc.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
    doc.add_paragraph("Заимствования")
    
    for i,url in enumerate(result_json['urls']):
        if url['url'] in url2color:
            doc.add_paragraph().add_run(f"({i}) {url['url']}").font.color.rgb = url2color[url['url']]
    
    doc.save(USER_FOLDER + 'full_text_report_result.docx')


def plagiarism_excepturl(request):
    exc_urls = request.POST.getlist('url')
    domains = [urlparse(url).hostname for url in exc_urls]
    # username = request.user.username
    with open(USER_FOLDER + 'file_name.pkl', 'rb') as f:
        file_name = pkl.load(f)
    with open(USER_FOLDER + file_name, 'r', encoding='UTF-8') as f:
        final_text = f.read()
    json_final = plagiarism_check(final_text, request.user, domains)
    with open(USER_FOLDER + 'json_final.pkl', 'wb') as f:
        pkl.dump(json_final, f)
    user = request.user
    generate_plag_report(user)
    if json_final is not None:
        if type(json_final.get('result_json', '{}')) == str:
            result_json = json.loads(json_final.get('result_json', '{}'))
        else:
            result_json = json_final.get('result_json', {})
        return HttpResponse(json.dumps({
            'uniqueness': json_final.get('unique'),
            'urls': result_json.get('urls', [])
        }))
    else:
        return HttpResponse("Error occurred while processing the file")

def download_file(request):
    username = request.user.username
    with open(USER_FOLDER + 'message_count.pkl','rb') as f:
        message_count = pkl.load(f)

    with open(USER_FOLDER + 'authors.pkl','rb') as f:
        authors = pkl.load(f)

    with open(USER_FOLDER + 'topic.pkl', 'rb') as f:
        topic = pkl.load(f)

    # generate spravka in pdf
    user = request.user
    generate_file_word(message_count, authors, user, topic)

    with open(USER_FOLDER + 'spravka.pdf', 'rb') as f:
        file_data = f.read()
    # Set the appropriate response headers to indicate a file download
    filename = name_the_file(username, 'spravka')
    quoted_filename = quote(filename.encode('utf-8'))
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{quoted_filename}.pdf'
    
    # Copy the contents of the file object to the response object
    response.write(file_data)
    
    return response

def download_report(request):
    username = request.user.username
    subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", USER_FOLDER + 'report.docx', "--outdir", USER_FOLDER])
    with open(USER_FOLDER + 'report.pdf','rb') as f:
        report_file = f.read()
    # Set the appropriate response headers to indicate a file download
    filename = name_the_file(username, 'avoid_report')
    quoted_filename = quote(filename.encode('utf-8'))
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{quoted_filename}.pdf'
    
    # Copy the contents of the file object to the response object
    response.write(report_file)
    
    # Close the file object
    # report_pdf.close()
    
    return response

def file_processed_avoid(request, response):
    data = json.loads(response)
    messages = []
    percentage = dict()
    error_mes = ''
    count_error = 0
    for match in data:
        percentage[match['message']] = match['errorLength']
    for match in data:
        # append the message to the list
        if match['message'] in messages:
            continue
        else:
            messages.append(match['message'])
        count_error += match['errorLength']
    error_mes = 'Кол-во ошибок:' + str(count_error)
    message = os.linesep.join(messages)
    percentage = json.dumps(percentage)
    # messages.append(error_mes)
    # message = os.linesep.join(messages)
    return render(request, 'avoid/file_processed_avoid.html', {"messages":message, "counter":error_mes, "percentage": percentage})

class BigFile:
    chnk_len = 150_000
    @staticmethod
    def create_text_chnks(final_text):
        chnk_len = BigFile.chnk_len
        texts = []
        for idx in range(0, len(final_text), chnk_len):
            texts.append(final_text[idx:idx + chnk_len])
        return texts
    
    @staticmethod
    async def process(final_text):
        texts = BigFile.create_text_chnks(final_text)
        tasks = []
        for text in texts:
            if len(text) < 300:
                break
            tasks.append(asyncio.create_task(BigFile.request_and_process(text)))
        json_final_responses = await asyncio.gather(*tasks)
        return json_final_responses


    @staticmethod
    async def request_and_process(text):
        url = 'https://api.text.ru/post'
        params = {
            'userkey': '31b045c683167166b5bae6c1014561cd',
            'text': text
        }
        response = requests.post(url, data=params)
        json_response = response.json()
        if response.status_code == 200 and json_response.get('error_code') is None:
            time.sleep(60)
            url = 'https://api.text.ru/post'
            params = {
                'userkey': '31b045c683167166b5bae6c1014561cd',
                'uid': json_response['text_uid'],
                'jsonvisible': 'detail'
            }
            response = requests.post(url, data=params)
            json_response_final = response.json()
            json_num = json_response_final.get('error_code')

            if (response.status_code == 200 and json_response_final.get('error_code') is None) :
                while json_num == 181:
                    time.sleep(10)
                    response = requests.post(url, data=params)
                    json_response_final = response.json()
                    json_num = json_response_final.get('error_code')
            else:
                json_response_final = {'error': json_response_final.get('error_desc', 'Возможно закончилась подписка')}
        
        else:
            json_response_final = {'error': json_response.get('error_desc', 'Error occurred while processing the file')}
        json_response_final['text_len'] = len(text)
        return json_response_final

    @staticmethod
    def is_bigfile(final_text):
        if len(final_text) >= BigFile.chnk_len:
            return True
        return False

    @staticmethod
    def plagiarism_check(final_text):
        json_responses = asyncio.run(BigFile.process(final_text))
        unique_final = 0
        offset = 0
        final_clear_text = ''
        final_result_json = {'urls':[]}
        urls = {}
        i=0
        for json_response in json_responses:
            if isinstance(json_response, Exception):
                print('got exception', json_response)
                return {'error':'ошибка сервера'}
            if 'error' in json_response:
                return {'error':json_response['error']}
            
            unique_final = unique_final + json_response['text_len']/len(final_text)*float(json_response['unique'])
            result_json = json.loads(json_response['result_json'])
            final_clear_text = final_clear_text + ' ' + result_json['clear_text']
        
            for url in result_json['urls']:
                if url['url'] not in urls:
                    words = list(map( lambda w: int(w)+offset, url['words'].split() ))
                    url['words'] = ' '.join([str(idx) for idx in words])
                    url['plagiat'] = url['plagiat']*len(words)
                    final_result_json['urls'].append(url)
                    urls[url['url']] = i
                    i+=1
                else:
                    words = list(map( lambda w: int(w)+offset, url['words'].split() ))
                    url_idx = urls[url['url']]
                    final_result_json['urls'][url_idx]['words'] = final_result_json['urls'][url_idx]['words'] + ' ' + ' '.join([str(idx) for idx in words])
                    final_result_json['urls'][url_idx]['plagiat'] += url['plagiat']*len(words)
                    
            offset = offset + len(result_json['clear_text'].split())
        
        final_text_len = len(final_clear_text.split())
        unique_final = round(unique_final,2)
        for url in final_result_json['urls']:
            url['plagiat'] = round(url['plagiat']/final_text_len,2)
        final_result_json['urls'].sort(reverse=True,key=lambda url:url['plagiat'])
        final_result_json['clear_text'] = final_clear_text
        final_result_json['unique'] = unique_final
        json_response_final={}
        json_response_final['result_json'] = final_result_json
        json_response_final['text_unique'] = str(unique_final)
        json_response_final['unique'] = str(unique_final)
        return json_response_final
